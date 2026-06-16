"""fridge.ai 발표 자료 워드 문서 생성기"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── 기본 폰트를 맑은 고딕으로 설정 ──────────────────────────────
for style_name in ('Normal', 'Default Paragraph Font'):
    try:
        s = doc.styles[style_name]
        s.font.name = 'Malgun Gothic'
        s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    except Exception:
        pass

# ── 페이지 여백 ──────────────────────────────────────────────────
section = doc.sections[0]
section.page_width   = Cm(21)
section.page_height  = Cm(29.7)
section.left_margin  = section.right_margin = Cm(2.5)
section.top_margin   = section.bottom_margin = Cm(2.5)

# ── 색상 ─────────────────────────────────────────────────────────
C_GREEN  = RGBColor(0x1A, 0x66, 0x40)
C_MID    = RGBColor(0x2F, 0x7D, 0x4F)
C_AMBER  = RGBColor(0xB8, 0x86, 0x0B)
C_GRAY   = RGBColor(0x55, 0x55, 0x55)
C_BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

FONT = 'Malgun Gothic'

# ── 공통 헬퍼 ────────────────────────────────────────────────────
def kr(run, size=10.5, bold=False, color=None):
    """한글 run에 Malgun Gothic 적용 (일반 + 동아시아 폰트 모두 설정)"""
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # 동아시아 폰트 설정 (핵심)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    FONT)
    rFonts.set(qn('w:hAnsi'),    FONT)
    rFonts.set(qn('w:eastAsia'), FONT)
    rFonts.set(qn('w:cs'),       FONT)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    return run

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_text(para, text, size=10.5, bold=False, color=None):
    run = para.add_run(text)
    return kr(run, size=size, bold=bold, color=color)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '2F7D4F')
    pb.append(bot)
    pPr.append(pb)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    add_text(p, text, size=16, bold=True, color=C_GREEN)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    add_text(p, text, size=13, bold=True, color=C_MID)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    add_text(p, text, size=11, bold=True, color=C_GRAY)

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_text(p, text, size=10.5, color=C_BLACK)

def body2(parts):
    """parts: [(text, bold), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    for text, bold in parts:
        add_text(p, text, size=10.5, bold=bold, color=C_BLACK)

def blt(text, bold_prefix=None, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.8)
    p.paragraph_format.space_after = Pt(3)
    # 수동 bullet
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl  = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), str(level))
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '1')
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)
    if bold_prefix:
        add_text(p, bold_prefix, size=10.5, bold=True, color=C_GREEN)
    add_text(p, text, size=10.5, color=C_BLACK)

def blt_simple(text, bold_prefix=None):
    """번호 없는 간단 불릿 (• 직접 삽입)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        add_text(p, bold_prefix + ' ', size=10.5, bold=True, color=C_GREEN)
    add_text(p, text, size=10.5, color=C_BLACK)

def make_header_row(table, headers, bg='1A6640'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(cell.paragraphs[0], h, size=10, bold=True, color=C_WHITE)

# ════════════════════════════════════════════════════════════════════
#  표지
# ════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
add_text(p, 'fridge.ai', size=32, bold=True, color=C_GREEN)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p2, '냉장고 재료 관리 & 환경 절감 서비스', size=15, color=C_MID)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p3, 'ERD 구조 · 데이터 설명 · 분석 방향성', size=12, color=C_GRAY)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p4, 'MANIFESTO v0.3  ·  2026-06-16', size=10, color=C_GRAY)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  PART 1. ERD 및 주요 내용
# ════════════════════════════════════════════════════════════════════
h1('PART 1. ERD 및 주요 내용')
divider()

h2('1-1. 서비스 개요')
body('fridge.ai는 3~4인 가정의 냉장고 재료를 스마트하게 관리하고, '
     '임박 재료를 활용한 레시피를 추천하여 식품 낭비와 탄소 배출을 줄이는 '
     '모바일 우선(Mobile-First) 웹 서비스입니다.')
doc.add_paragraph()
blt_simple('유통기한 D-day 기반 재고 관리 (임박 / 주의 / 신선 3단계)')
blt_simple('보유 재료 매칭 레시피 추천 (스코어링 엔진)')
blt_simple('월간 탄소 절감량 · 절약 식비 리포트')
blt_simple('신규 재료 요청 큐 (F10 — 관리자 검토 후 마스터 반영)')

doc.add_paragraph()
h2('1-2. ERD 구조 개요 (8개 테이블)')
body('총 8개 테이블로 구성되며, 재료 마스터 · 재고 · 요리 DB · 소진 기록의 4개 영역으로 구분됩니다.')

tbl1 = doc.add_table(rows=9, cols=3)
tbl1.style = 'Table Grid'
make_header_row(tbl1, ['테이블', '영역', '역할'])
tbl1.columns[0].width = Cm(4.5)
tbl1.columns[1].width = Cm(3.5)
tbl1.columns[2].width = Cm(9.0)

rows_data = [
    ('users',                '사용자',     '계정 정보 (이메일, 비밀번호 해시)'),
    ('categories',           '재료 마스터', '재료 카테고리 10종 (채소 · 육류 · 해산물 등)'),
    ('ingredients',          '재료 마스터', '재료 마스터 131종 — 단위 · 탄소발자국 · 무게 포함'),
    ('user_inventory',       '재고',        '개인 냉장고 재고 — 수량 · 유통기한 · 상태 관리'),
    ('recipes',              '요리 DB',     '요리/메뉴 마스터 45종 (국/찌개 · 볶음 · 구이 등)'),
    ('recipe_ingredients',   '요리 DB',     '요리 ↔ 재료 매핑 (필수/선택 구분)'),
    ('consumption_logs',     '소진 기록',   '조리(cooked) / 폐기(wasted) 이력 — 탄소 계산 기반'),
    ('ingredient_requests',  '요청 큐',     'F10 신규 재료 요청 — 관리자 검토 후 마스터 반영'),
]
for i, (name, area, desc) in enumerate(rows_data):
    row = tbl1.rows[i + 1]
    color = 'FFFFFF' if i % 2 == 0 else 'F4FAF6'
    set_cell_bg(row.cells[0], color)
    set_cell_bg(row.cells[1], color)
    set_cell_bg(row.cells[2], color)
    add_text(row.cells[0].paragraphs[0], name, size=9.5, bold=True, color=C_GREEN)
    add_text(row.cells[1].paragraphs[0], area, size=9.5, color=C_BLACK)
    add_text(row.cells[2].paragraphs[0], desc, size=9.5, color=C_BLACK)

doc.add_paragraph()
h2('1-3. 핵심 데이터 상세 설명')

h3('① ingredients — 재료 마스터')
body('총 131종의 재료를 10개 카테고리로 분류합니다. '
     '각 재료는 기본 단위, 단위당 무게(g), 탄소발자국(kg CO2eq / 100g) 값을 보유합니다.')

tbl2 = doc.add_table(rows=5, cols=2)
tbl2.style = 'Table Grid'
data2 = [
    ('name',              '재료명 (고유값, 자동완성 · 스코어링 기준)'),
    ('default_unit',      '기본 단위 — 개 / g / mL / 팩 / 봉'),
    ('weight_per_unit_g', '1단위 기본 무게 (예: 달걀 1개=60g, 감자 1개=150g)'),
    ('carbon_per_100g',   '탄소발자국 (kg CO2eq / 100g) — 농진청 · 문헌 기반'),
    ('카테고리 10종',      '채소 · 육류 · 해산물 · 달걀/유제품 · 두부/콩류 · 버섯류 · 곡류/면류 · 과일 · 조미료/양념 · 기타'),
]
for i, (col, desc) in enumerate(data2):
    set_cell_bg(tbl2.rows[i].cells[0], 'D4EDDA')
    add_text(tbl2.rows[i].cells[0].paragraphs[0], col, size=9.5, bold=True, color=C_GREEN)
    add_text(tbl2.rows[i].cells[1].paragraphs[0], desc, size=9.5, color=C_BLACK)

doc.add_paragraph()
h3('② recipes + recipe_ingredients — 요리/메뉴 DB')
body('총 45종 요리와 재료 매핑을 구축합니다. '
     '각 요리는 필수 재료 목록을 보유하며, 스코어링 엔진이 보유 재료와 매칭합니다.')

tbl3 = doc.add_table(rows=3, cols=2)
tbl3.style = 'Table Grid'
data3 = [
    ('카테고리 8종',  '국/찌개 · 볶음 · 전/부침 · 구이 · 조림 · 밥/면 · 무침 · 기타'),
    ('is_required',  '필수 재료(TRUE) / 선택 재료(FALSE) → 스코어링 정밀도 향상'),
    ('source_url',   'YouTube/블로그 링크 캐싱용 (Phase 2, MVP는 검색 링크 자동 생성)'),
]
for i, (col, desc) in enumerate(data3):
    set_cell_bg(tbl3.rows[i].cells[0], 'FFF3CD')
    add_text(tbl3.rows[i].cells[0].paragraphs[0], col, size=9.5, bold=True, color=C_AMBER)
    add_text(tbl3.rows[i].cells[1].paragraphs[0], desc, size=9.5, color=C_BLACK)

doc.add_paragraph()
h3('③ user_inventory — 냉장고 재고')
body('사용자가 등록한 재료의 수량 · 유통기한 · 상태를 관리합니다. expires_at 기준으로 D-day를 계산합니다.')
blt_simple('임박 (D 2일 이하): 빨간 경고')
blt_simple('주의 (D 3~5일): 노란 경고')
blt_simple('신선 (D 6일+): 정상')
blt_simple('status: active → used(조리) / wasted(폐기) 로 전환')

doc.add_paragraph()
h3('④ consumption_logs — 소진/폐기 기록')
body('모든 소진 이력을 저장하며, weight_g × carbon_per_100g / 100 공식으로 탄소량을 역산합니다.')
blt_simple('type = cooked: 조리 완료 → 탄소 절감으로 산정')
blt_simple('type = wasted: 폐기 → 탄소 손실로 산정')
blt_simple('recipe_id (FK?): 참고한 레시피 연결 (NULL 허용)')

doc.add_paragraph()
h3('⑤ ingredient_requests — 신규 재료 요청 큐 (F10)')
body('사용자가 DB에 없는 재료를 요청하면 pending 상태로 큐에 쌓입니다. '
     '관리자가 검토 후 approved 처리하면 ingredients 마스터에 수동 추가됩니다.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  PART 2. 데이터 분석 방향성
# ════════════════════════════════════════════════════════════════════
h1('PART 2. 데이터 분석 방향성')
divider()

body('fridge.ai가 수집하는 데이터를 활용하여 아래 4가지 분석을 수행합니다. '
     '각 분석은 "추천(Recommend)" 또는 "진단(Diagnose)" 형태로 사용자에게 제공됩니다.')

doc.add_paragraph()

# 분석 1
h2('분석 1.  임박 재료 기반 레시피 추천  (Recommend)')
body2([
    ('핵심 질문: ', True),
    ('"지금 냉장고에 있는 재료로, 가장 빨리 소진해야 할 재료를 활용할 수 있는 요리는 무엇인가?"', False),
])

h3('활용 데이터')
blt_simple('user_inventory.expires_at')
blt_simple('recipe_ingredients.ingredient_id (is_required = TRUE)')
blt_simple('ingredients.name (재료 매칭 키)')

h3('스코어링 공식')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.8)
p.paragraph_format.space_after = Pt(6)
add_text(p, 'Score  =  40 x (임박 재료 사용 비율)  +  30 x (보유 재료 매칭율)  -  10 x (부족 재료 페널티)',
         size=10.5, bold=True, color=C_GREEN)

blt_simple('임박 재료(D 2일 이하)를 많이 쓸수록 가중치 40점 부여 → 버리기 전에 소진 유도')
blt_simple('보유 재료 매칭율이 높을수록 30점 추가 → 추가 구매 없이 조리 가능한 레시피 우선')
blt_simple('보유하지 않은 재료가 많으면 10점 차감 → 재료 부족 레시피 후순위')

h3('산출물')
blt_simple('홈 화면 "오늘의 추천" 카드 상위 3종')
blt_simple('추천 탭 전체 레시피 스코어 순 정렬 (최대 12종)')
blt_simple('각 레시피 카드에 재료 태그 (임박 / 보유 / 미보유) 표시')

doc.add_paragraph()

# 분석 2
h2('분석 2.  월간 환경 영향 진단  (Diagnose)')
body2([
    ('핵심 질문: ', True),
    ('"이번 달 나의 냉장고 관리로 탄소를 얼마나 절감했는가?"', False),
])

h3('활용 데이터')
blt_simple('consumption_logs.weight_g × ingredients.carbon_per_100g / 100')
blt_simple('consumption_logs.type = cooked / wasted 구분')

h3('분석 지표')
tbl4 = doc.add_table(rows=5, cols=3)
tbl4.style = 'Table Grid'
make_header_row(tbl4, ['지표', '계산 방식', '의미'])
data4 = [
    ('탄소 절감량 (kg CO2)',  'SUM(weight_g x carbon / 100) — type=cooked', '조리 완료로 절약된 탄소량'),
    ('탄소 손실량 (kg CO2)',  'SUM(weight_g x carbon / 100) — type=wasted', '폐기로 낭비된 탄소량'),
    ('절약 식비 (원)',        'weight_g / 1000 x 재료 시세 (추정)',           '폐기 방지로 아낀 금액'),
    ('소진율 (%)',            'cooked / (cooked + wasted) x 100',            '전체 소진 중 조리 비율'),
]
for i, (metric, formula, meaning) in enumerate(data4):
    row = tbl4.rows[i + 1]
    color = 'F4FAF6' if i % 2 == 0 else 'FFFFFF'
    for j, (val, c, b) in enumerate([
        (metric, C_GREEN, True), (formula, C_BLACK, False), (meaning, C_BLACK, False)
    ]):
        set_cell_bg(row.cells[j], color)
        add_text(row.cells[j].paragraphs[0], val, size=9, bold=b, color=c)

doc.add_paragraph()

# 분석 3
h2('분석 3.  식품 폐기 패턴 진단  (Diagnose)')
body2([
    ('핵심 질문: ', True),
    ('"어떤 재료를 언제 가장 많이 버리는가? 무엇을 개선할 수 있는가?"', False),
])

h3('활용 데이터')
blt_simple('consumption_logs (type=wasted) × ingredients × categories')

h3('분석 항목')
blt_simple('카테고리별 폐기율: 어느 카테고리가 낭비가 심한가', bold_prefix='①')
blt_simple('재료별 평균 잔여 D-day: 얼마나 일찍 버리는가 (expires_at - logged_at)', bold_prefix='②')
blt_simple('월별 추이: 폐기량이 줄어들고 있는가', bold_prefix='③')
blt_simple('폐기 시점 분석: 어떤 요일/시간대에 폐기가 집중되는가', bold_prefix='④')

h3('산출물')
blt_simple('"이번 달 가장 많이 버린 재료 TOP 3" 알림')
blt_simple('"채소류 폐기율이 전월 대비 -15% 개선됐어요" 피드백 메시지')
blt_simple('리포트 탭 주간 소진 추이 바 차트')

doc.add_paragraph()

# 분석 4
h2('분석 4.  재료 구매 패턴 추천  (Recommend)')
body2([
    ('핵심 질문: ', True),
    ('"자주 소진되는 재료는 무엇이며, 언제 다시 구매해야 하는가?"', False),
])

h3('활용 데이터')
blt_simple('consumption_logs (type=cooked) — 재료별 조리 빈도')
blt_simple('user_inventory — 현재 재고 수량 · 유통기한')
blt_simple('ingredient_requests — 사용자 요청 빈도')

h3('분석 항목')
blt_simple('자주 조리에 활용된 재료 상위 N개 → "자주 사는 재료" 빠른 추가 버튼 학습', bold_prefix='①')
blt_simple('재고 소진 주기 계산: 평균 몇 일 만에 소진되는가', bold_prefix='②')
blt_simple('재고 소진 예측: expires_at 이전에 소진 가능한지 판단', bold_prefix='③')
blt_simple('요청 큐 빈도 분석: 어떤 재료가 자주 요청되는가 → 마스터 추가 우선순위', bold_prefix='④')

h3('산출물')
blt_simple('"달걀이 3일 내 소진 예정입니다. 미리 구매하세요" 알림 (Phase 2)')
blt_simple('"자주 사는 재료" 추천 버튼 개인화 (현재는 고정 9종)')
blt_simple('요청 빈도 Top 10 → 관리자 마스터 추가 우선순위 결정')

doc.add_paragraph()
divider()
doc.add_paragraph()

# 요약 테이블
h2('분석 방향성 요약')
tbl5 = doc.add_table(rows=5, cols=4)
tbl5.style = 'Table Grid'
make_header_row(tbl5, ['분석', '유형', '핵심 데이터', '사용자 가치'])
summary = [
    ('임박 재료 레시피 추천', 'Recommend', 'user_inventory + recipe_ingredients', '음식 낭비 없이 오늘 저녁 메뉴 결정'),
    ('월간 환경 영향 진단',   'Diagnose',  'consumption_logs + ingredients.carbon', '탄소 절감량 · 절약 식비 수치 확인'),
    ('식품 폐기 패턴 진단',   'Diagnose',  'consumption_logs (wasted) + categories', '어떤 재료를 자주 버리는지 인식'),
    ('구매 패턴 추천',        'Recommend', 'consumption_logs (cooked) + ingredient_requests', '언제 무엇을 사야 할지 예측'),
]
for i, (name, typ, data, value) in enumerate(summary):
    row = tbl5.rows[i + 1]
    color = 'F4FAF6' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        set_cell_bg(cell, color)
    add_text(row.cells[0].paragraphs[0], name, size=9.5, bold=True, color=C_GREEN)
    typ_color = C_MID if typ == 'Recommend' else C_AMBER
    add_text(row.cells[1].paragraphs[0], typ, size=9.5, bold=True, color=typ_color)
    add_text(row.cells[2].paragraphs[0], data, size=9, color=C_BLACK)
    add_text(row.cells[3].paragraphs[0], value, size=9, color=C_BLACK)

# ── 저장 ─────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fridge_ai_presentation.docx')
doc.save(out)
print(f'saved: {out}')
